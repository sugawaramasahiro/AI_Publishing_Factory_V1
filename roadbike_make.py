from docx import Document
from ebooklib import epub
import os, html, datetime

docx_path   = r"C:\AI_Publishing_Factory_V1\ロードバイク30km本_TMP.docx"
cover_path  = r"C:\AI_Publishing_Factory_V1\assets\cover_sample.jpg"
output_docx = r"C:\AI_Publishing_Factory_V1\roadbike_fixed.docx"
output_epub = r"C:\AI_Publishing_Factory_V1\roadbike_book.epub"

# ---------- Word 整形 ----------
doc = Document(docx_path)
tp  = doc.add_paragraph(); tp.alignment = 1
tp.add_run("ロードバイク初心者が\\n30km/h巡航する方法").bold = True
tp.add_run("\\n\\nMP").bold = True
doc.add_page_break()

doc.add_paragraph("目次", style="Heading 1")
doc.add_page_break()

doc.add_page_break()
p = doc.add_paragraph()
p.add_run("奥付").bold = True
p.add_run(f"\\n発行日：{datetime.date(2025,5,29)}")
p.add_run("\\n著者：MP\\n出版社：MP\\n©2025 MP. All rights reserved.")
doc.save(output_docx)
print("✅ Word 完成 →", output_docx)

# ---------- EPUB 生成 ----------
book = epub.EpubBook()
book.set_title("ロードバイク初心者が30km/h巡航する方法")
book.set_language("ja"); book.add_author("MP")
if os.path.exists(cover_path):
    book.set_cover("cover.jpg", open(cover_path,'rb').read())

chapters, cur, body = [], "はじめに", ""
for para in Document(output_docx).paragraphs:
    t = para.text.strip()
    if not t: continue
    if t.startswith("第") and "章" in t:
        if body:
            h = f"<h1>{cur}</h1><p>{html.escape(body).replace(chr(10),'<br/>')}</p>"
            sec = epub.EpubHtml(title=cur,file_name=f"{cur}.xhtml",lang="ja"); sec.content=h
            book.add_item(sec); chapters.append(sec)
        cur, body = t, ""
    else:
        body += t + "\\n"
if body:
    h = f"<h1>{cur}</h1><p>{html.escape(body).replace(chr(10),'<br/>')}</p>"
    sec = epub.EpubHtml(title=cur,file_name=f"{cur}.xhtml",lang="ja"); sec.content=h
    book.add_item(sec); chapters.append(sec)

book.toc = chapters
book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav())
epub.write_epub(output_epub, book)
print("✅ EPUB 完成 →", output_epub)
