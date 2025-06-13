from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def export_to_word(title, author, chapters, output_path="output"):
    doc = Document()

    # 表紙
    doc.add_paragraph(title).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(author).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # 本文
    for chapter_title, content in chapters.items():
        doc.add_heading(chapter_title, level=1)
        for paragraph in content.strip().split("\n"):
            doc.add_paragraph(paragraph.strip())
        doc.add_page_break()

    # 奥付
    doc.add_heading("奥付", level=1)
    doc.add_paragraph(f"著者名：{author}")
    doc.add_paragraph("出版社：MP出版（AI出版所）")
    doc.add_paragraph("発行日：2025年")
    doc.add_paragraph("本書の著作権は著者に帰属します。")
    doc.add_page_break()

    os.makedirs(output_path, exist_ok=True)
    filename = os.path.join(output_path, f"{title}.docx")
    doc.save(filename)

    print(f"✅ Wordファイルを書き出しました：{filename}")
    return filename

# CLI用
if __name__ == "__main__":
    sample_chapters = {
        "はじめに": "これはテストです。\n複数行対応しています。",
        "第1章 睡眠の重要性": "睡眠は体と心にとってとても大切です。\n良い睡眠をとりましょう。",
        "おわりに": "ご覧いただきありがとうございました。"
    }
    export_to_word("AI出版テンプレ本", "MP", sample_chapters)
