import docx, datetime, os

doc = docx.Document(r"'$doc'")
# 0. タイトルページを先頭に追加
title_page = doc.add_paragraph()
title_page.alignment = 1           # 中央
title_page.add_run("ビジネスパーソンのための\n「よく眠れる習慣」").bold = True
title_page.add_run("\n\nMP").bold = True
doc.add_page_break()

# 1. 目次プレースホルダ
toc = doc.add_paragraph()
toc.add_run("目次").bold = True
doc.add_page_break()

# 2. 奥付ページを末尾に追加
doc.add_page_break()
p = doc.add_paragraph()
p.add_run("奥付").bold = True
p.add_run(f"\n発行日：{datetime.date(2025,5,29)}")
p.add_run("\n著者：MP\n出版社：MP\n©2025 MP. All rights reserved.")
doc.save(r"'$doc_fixed'")
print("✅ Word 整形完了 →", r"'$doc_fixed'")
