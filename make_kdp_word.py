from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date

doc = Document()

# 表紙
title = doc.add_paragraph()
title.alignment = 1
title.add_run("ロードバイク初心者が\n30km/h巡航する方法").bold = True
title.add_run("\n\n著者：MP").italic = True
doc.add_page_break()

# 自動目次
p = doc.add_paragraph()
r = p.add_run()
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
r._r.append(fld)
doc.add_page_break()

def add_heading(text, level): doc.add_paragraph(text, style=f"Heading {level}")
def add_body(text): [doc.add_paragraph(l) for l in text.strip().split("\n")]

# まえがき
add_heading("まえがき", 1)
add_body("""
ロードバイクに乗り始めた頃、誰もが憧れる数字「30km/h」。
でも実際に維持するのは難しい。
本書では、その壁を超えるための知識・技術・心構えを、
初心者目線で解説しています。
""")
doc.add_page_break()

# 第1章
add_heading("第1章：バイクの選び方", 1)
add_heading("1-1 フレームサイズとフィッティング", 2)
add_body("サイズが合わないと巡航どころか疲労がたまる。")
add_heading("1-2 初心者向け装備の選び方", 2)
add_body("最初は軽量より快適性重視で選ぶのが正解。")
doc.add_page_break()

# 第2章
add_heading("第2章：体力と技術の基礎", 1)
add_heading("2-1 心拍ゾーンとトレーニング", 2)
add_body("Z2ゾーンで60分走る体力が巡航の鍵。")
add_heading("2-2 ケイデンスとフォーム", 2)
add_body("90rpmを意識。脱力フォームとペダリング精度。")
doc.add_page_break()

# あとがき
add_heading("あとがき", 1)
add_body("30km/hは“目標”ではなく“通過点”。あなたもきっと超えられる。")

# プロフィール
add_heading("著者プロフィール", 1)
add_body("MP（エムピー）：市民サイクリスト。本書が初出版。")

# 奥付
doc.add_page_break()
add_heading("奥付", 1)
add_body(f"発行日：{date.today().isoformat()}\n著者：MP\n出版社：MP\n©{date.today().year} MP. All rights reserved.")

doc.save(r"C:\AI_Publishing_Factory_V1\book_ready.docx")
print("✅ Word（KDP提出用）を書き出しました →", r"C:\AI_Publishing_Factory_V1\book_ready.docx")
