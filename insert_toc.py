from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os, datetime

doc = Document(r"C:\AI_Publishing_Factory_V1\ロードバイク30km本_TMP.docx")

# タイトルページ（中央寄せ）
title = doc.add_paragraph()
title.alignment = 1
title.add_run("ロードバイク初心者が\\n30km/h巡航する方法").bold = True
title.add_run("\\n\\nMP").bold = True
doc.add_page_break()

# ▶ 本物の目次フィールドを挿入
p = doc.add_paragraph()
r = p.add_run()
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
r._r.append(fld)
doc.add_page_break()

# 奥付ページ
doc.add_page_break()
footer = doc.add_paragraph()
footer.add_run("奥付").bold = True
footer.add_run(f"\n発行日：{datetime.date(2025,5,29)}")
footer.add_run("\n著者：MP\n出版社：MP\n©2025 MP. All rights reserved.")

# 保存
doc.save(r"C:\AI_Publishing_Factory_V1\roadbike_with_toc.docx")
print("✅ 自動目次入り Word 完成 →", r"C:\AI_Publishing_Factory_V1\roadbike_with_toc.docx")
