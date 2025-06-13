from openai import OpenAI
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
import os

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def ask_gpt(prompt):
    res = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return res.choices[0].message.content.strip()

# --- Word 作成 ---
doc = Document()
title = "ロードバイク初心者が30km/h巡航する方法"
author = "MP"

# 表紙
tp = doc.add_paragraph(); tp.alignment = 1
tp.add_run(title).bold = True
tp.add_run(f"\n\n著者：{author}").italic = True
doc.add_page_break()

# 自動目次フィールド
p = doc.add_paragraph()
r = p.add_run()
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
r._r.append(fld)
doc.add_page_break()

# ヘルパー関数
def add_heading(text, level): doc.add_paragraph(text, style=f"Heading {level}")
def add_body(text): [doc.add_paragraph(line) for line in text.strip().split("\\n")]

# まえがき
add_heading("まえがき", 1)
add_body(ask_gpt("ロードバイク初心者が30km/h巡航したい理由と悩みを、400字程度でまえがき風に書いてください。"))

# 本文（2章分）
chapters = [
    ("第1章：バイクの選び方", [
        "初心者がやりがちなバイク選びの失敗と正解を、400字で書いてください。",
        "初心者におすすめの装備（タイヤ・ギア・ヘルメットなど）について400字で説明してください。"
    ]),
    ("第2章：体力と技術の基礎", [
        "Z2トレーニングと心拍ゾーンを説明し、なぜ重要かを初心者向けに解説してください（400字）。",
        "ケイデンスとフォームを改善するメリットを初心者向けに400字で書いてください。"
    ])
]

for ch_title, prompts in chapters:
    add_heading(ch_title, 1)
    for i, p in enumerate(prompts, 1):
        add_heading(f"{ch_title.split('：')[0]}-{i}", 2)
        add_body(ask_gpt(p))
    doc.add_page_break()

# あとがき
add_heading("あとがき", 1)
add_body(ask_gpt("30km/h巡航を達成した達成感と、次の目標に向かう読者への応援メッセージを300字で書いてください。"))

# 著者プロフィール
add_heading("著者プロフィール", 1)
add_body("MP（エムピー）：市民サイクリスト。本書が初出版。トレーニングと電子出版に情熱を注ぐ。")

# 奥付
doc.add_page_break()
add_heading("奥付", 1)
add_body(f"発行日：{date.today().isoformat()}\n著者：MP\n発行者：MP\n©{date.today().year} MP. All rights reserved.")

# 書き出し
doc.save(r"C:\AI_Publishing_Factory_V1\book_api_ready.docx")
print("✅ Word（最新版API使用）完成 →", r"C:\AI_Publishing_Factory_V1\book_api_ready.docx")
