from docx import Document
import os
from datetime import datetime

def create_word_doc():
    book = {
        "title": "ローカル出版所V100デモ",
        "author": "MP",
        "chapters": [
            {"title": "第1章：はじめに", "content": "この本は完全ローカルで作成しています。"},
            {"title": "第2章：構成", "content": "全処理をローカルPCだけで完結させます。"},
            {"title": "第3章：安全性", "content": "外部にデータは一切送信しません。"}
        ]
    }

    doc = Document()
    doc.add_heading(book["title"], 0)
    doc.add_paragraph(book["author"])
    doc.add_page_break()

    doc.add_heading("目次", level=1)
    for chapter in book["chapters"]:
        doc.add_paragraph(chapter["title"])
    doc.add_page_break()

    for chapter in book["chapters"]:
        doc.add_heading(chapter["title"], level=1)
        doc.add_paragraph(chapter["content"])
        doc.add_page_break()

    os.makedirs("output", exist_ok=True)
    # バージョン付与用タイムスタンプ
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    fname = f"output/local_demo_{ts}.docx"
    doc.save(fname)
    # 最新ファイルも同時保存
    doc.save("output/local_demo.docx")
    print("✅ Word原稿生成・バックアップ完了！", fname)

if __name__ == "__main__":
    create_word_doc()
